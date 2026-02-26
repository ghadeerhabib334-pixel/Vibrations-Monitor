from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('Vibration & Temperature Monitoring Dashboard', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Real-Time Industrial IoT Monitoring System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(14)

doc.add_paragraph()

# Project Overview
doc.add_heading('Project Overview', level=1)
doc.add_paragraph(
    'A web-based real-time monitoring dashboard designed for industrial vibration and temperature '
    'analysis. The system connects to Erbessd Phantom sensors via MQTT protocol, providing instant '
    'visibility into machine health and enabling predictive maintenance decisions.'
)

# Key Features
doc.add_heading('Key Features', level=1)

features = [
    ('Real-Time Vibration Monitoring', 'Displays tri-axial (X, Y, Z) vibration RMS values with ISO 10816 severity classification (Good, Satisfactory, Unsatisfactory, Danger).'),
    ('Temperature Tracking', 'Continuous temperature monitoring with min/max tracking and visual gauge representation.'),
    ('System Health Indicators', 'Battery voltage monitoring and wireless signal strength (RSSI) visualization.'),
    ('Historical Data Charts', 'Interactive trend charts for vibration and temperature data using Chart.js.'),
    ('MQTT Connectivity', 'Configurable MQTT broker connection with support for secure WebSocket (WSS) protocol and topic wildcards.'),
    ('Responsive Design', 'Modern glass-morphism UI with dark/light theme support, optimized for desktop and mobile devices.'),
]

for feature_title, feature_desc in features:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(feature_title + ': ').bold = True
    p.add_run(feature_desc)

# Technologies Used
doc.add_heading('Technologies Used', level=1)

tech_paragraph = doc.add_paragraph()
technologies = ['HTML5', 'CSS3 (Tailwind CSS)', 'JavaScript', 'MQTT.js', 'Chart.js', 'WebSocket']
tech_paragraph.add_run(', '.join(technologies))

# Use Case
doc.add_heading('Use Case', level=1)
doc.add_paragraph(
    'This dashboard is deployed for monitoring industrial machinery equipped with Erbessd Phantom '
    'EPH-V11E wireless vibration sensors. It enables maintenance teams to identify potential equipment '
    'failures before they occur, reducing downtime and maintenance costs through condition-based monitoring.'
)

# My Role
doc.add_heading('My Role', level=1)
doc.add_paragraph(
    'Full-stack development including UI/UX design, front-end implementation, MQTT integration, '
    'and real-time data visualization. Implemented ISO 10816 vibration severity standards for '
    'industry-compliant machine health assessment.'
)

# Save
doc.save('C:/Users/user/Desktop/Vibration sensor/Project_Description.docx')
print('Document created successfully: Project_Description.docx')
